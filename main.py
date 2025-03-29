import streamlit as st
from st_on_hover_tabs import on_hover_tabs
import streamlit.components.v1 as components
from typing import List, Dict
from langchain_groq import ChatGroq
from langchain.chains import LLMChain, ConversationChain
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
import json
from docx import Document
import os,re
import subprocess
import tempfile
from streamlit_lottie import st_lottie

hide_menu_style = """
        <style>
        MainMenu {visibility: hidden;}
        footer {visibility: hidden; }
        </style>
        """
st.markdown(hide_menu_style, unsafe_allow_html=True)

try:
    with open("style.css") as f:
        st.markdown('<style>' + f.read() + '</style>', unsafe_allow_html=True)
except Exception:
    pass

dark_purple_theme = """
<style>
    /* Main theme colors */
    :root {
        --primary-color: #7B2CBF;
        --background-color: black;
        --secondary-bg: #2D2D2D;
        --text-color: #FFFFFF;
        --accent-color: #9D4EDD;
    }

    /* Main background */
    .stApp {
        background-color: var(--background-color);
        color: var(--text-color);
    }

    /* Headers */
    h1, h2, h3 {
        color: var(--accent-color) !important;
    }

    /* Buttons */
    .stButton>button {
        background-color: var(--primary-color);
        color: white;
        border: none;
        border-radius: 4px;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: var(--accent-color);
        transform: translateY(-2px);
    }

    /* Text areas and inputs */
    .stTextArea>div>div>textarea {
        background-color: var(--secondary-bg);
        color: var(--text-color);
        border: 1px solid var(--primary-color);
    }

    /* File uploader */
    .stFileUploader {
        background-color: var(--secondary-bg);
        border: 1px dashed var(--primary-color);
        border-radius: 4px;
    }

    /* Progress bar */
    .stProgress > div > div > div > div {
        background-color: var(--primary-color);
    }

    /* Chat messages */
    .stChatMessage {
        background-color: var(--secondary-bg);
        border-radius: 8px;
        padding: 10px;
        margin: 5px 0;
    }
    }
</style>
"""

def load_lottie_urls():
    urls = {
        "shield": "https://assets5.lottiefiles.com/packages/lf20_yom6uvgj.json",
        "analysis": "https://assets8.lottiefiles.com/packages/lf20_qmfs6c3i.json",
        "chat": "https://assets8.lottiefiles.com/packages/lf20_2LdLki.json",
        "security": "https://assets8.lottiefiles.com/packages/lf20_oyi9a28g.json"
    }
    return {key: url for key, url in urls.items()}

def about_section():
    animations = load_lottie_urls()
    
    # Header Section
    st.title("GAMKERS Security Analysis Suite")
    st.write("---")

    # First Container: Main Feature
    with st.container():
        left_col, right_col = st.columns(2)
        with left_col:
            st.header("Advanced Security Analysis")
            st.write("""
            Our AI-powered security suite provides deep code inspection and vulnerability detection.
            Using state-of-the-art machine learning models, we analyze your code for potential security
            threats and malicious patterns.
            """)
            st.button("🚀 Try Analysis Now", key="try_analysis")
        with right_col:
            st_lottie(animations["analysis"], height=300, key="analysis_anim")

    st.write("---")

    # Second Container: Chat Expert
    with st.container():
        left_col, right_col = st.columns(2)
        with left_col:
            st_lottie(animations["chat"], height=300, key="chat_anim")
        with right_col:
            st.header("Real-time Expert Consultation: GAMKERSGPT")
            st.write("""
            Get instant security advice from our AI expert system. Ask questions about:
            - Code vulnerabilities
            - Security best practices
            - Threat detection
            - Risk mitigation strategies
            """)
            st.button("💬 Start Chat", key="start_chat")

    st.write("---")

    # Third Container: Security Features
    with st.container():
        left_col, right_col = st.columns(2)
        with left_col:
            st.header("Comprehensive Security Suite")
            col1, col2 = st.columns(2)
            with col1:
                st.write("🔍 Code Obfuscation")
                st.write("🌐 Network Analysis")
                st.write("📂 File Operations")
            with col2:
                st.write("⚠️ Payload Detection")
                st.write("🔐 API Security")
                st.write("🛡️ Anti-Analysis")
        with right_col:
            st_lottie(animations["security"], height=300, key="security_anim")

    st.write("---")

    # Fourth Container: Stats and Info
    with st.container():
        st.header("Security Insights")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Languages Supported", "7+")
        with col2:
            st.metric("Security Rules", "1000+")
        with col3:
            st.metric("Analysis Speed", "<2 min")

    st.write("---")

class SecurityAnalysisApp:
    def __init__(self):
        self.chat_model = ChatGroq(
            groq_api_key="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",
            model_name="qwen-qwq-32b",
            temperature=0.7,
            max_tokens=None
        )
        
        self.analysis_template = PromptTemplate(
            input_variables=["code_chunk"],
            template="""Analyze the following code for malicious indicators and security concerns.

Code to analyze:
{code_chunk}

Provide a comprehensive analysis with a brief summary of findings and details for each category. Return your analysis in this exact JSON structure:
{{
    "summary": ["Brief overview of key findings and potential security implications"],
    "sections": {{
        "code_obfuscation_techniques": {{
            "findings": [],
            "description": "Brief explanation of identified obfuscation techniques and their implications"
        }},
        "suspicious_api_calls": {{
            "findings": [],
            "description": "Overview of concerning API calls and their potential security impact"
        }},
        "anti_analysis_mechanisms": {{
            "findings": [],
            "description": "Summary of detected anti-analysis features and their significance"
        }},
        "network_communication_patterns": {{
            "findings": [],
            "description": "Analysis of network-related code patterns and security concerns"
        }},
        "file_system_operations": {{
            "findings": [],
            "description": "Evaluation of file system interactions and associated risks"
        }},
        "potential_payload_analysis": {{
            "findings": [],
            "description": "Assessment of potential malicious payloads and their characteristics"
        }}
    }}
}}

Requirements:
1. Ensure each field is populated with meaningful content
2. Include a clear summary of overall findings
3. Provide a brief description for each section
4. List specific findings as bullet points in the findings arrays
5. Use "None identified" in findings array if no indicators are found
6. Keep descriptions concise and focused on security implications

Your response should be ONLY the JSON object with no additional text."""
        )
        
        self.binary_analysis_template = PromptTemplate(
            input_variables=["strings_chunk"],
            template="""Analyze the following strings extracted from a binary file for malicious indicators and security concerns.

Extracted strings:
{strings_chunk}

Provide a comprehensive analysis with a brief summary of findings and details for each category. Return your analysis in this exact JSON structure:
{{
    "summary": ["Brief overview of key findings and potential security implications"],
    "sections": {{
        "suspicious_strings": {{
            "findings": [],
            "description": "Brief explanation of identified suspicious strings and their implications"
        }},
        "command_and_control_indicators": {{
            "findings": [],
            "description": "Overview of potential C2 indicators like URLs, IPs, or domain patterns"
        }},
        "anti_analysis_indicators": {{
            "findings": [],
            "description": "Summary of strings suggesting anti-analysis capabilities"
        }},
        "network_related_strings": {{
            "findings": [],
            "description": "Analysis of network-related strings and security concerns"
        }},
        "file_system_indicators": {{
            "findings": [],
            "description": "Evaluation of file system related strings and associated risks"
        }},
        "potential_malware_functionality": {{
            "findings": [],
            "description": "Assessment of strings indicating malicious functionality"
        }}
    }}
}}

Requirements:
1. Ensure each field is populated with meaningful content
2. Include a clear summary of overall findings
3. Provide a brief description for each section
4. List specific findings as bullet points in the findings arrays
5. Use "None identified" in findings array if no indicators are found
6. Keep descriptions concise and focused on security implications

Your response should be ONLY the JSON object with no additional text."""
        )
        
        self.analysis_chain = LLMChain(
            llm=self.chat_model,
            prompt=self.analysis_template,
            verbose=True
        )
        
        self.binary_analysis_chain = LLMChain(
            llm=self.chat_model,
            prompt=self.binary_analysis_template,
            verbose=True
        )
        
        self.chat_memory = ConversationBufferMemory()
        self.conversation = ConversationChain(
            llm=self.chat_model,
            memory=self.chat_memory,
            verbose=True
        )

    def clean_json_response(self, response: str) -> str:
        try:
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end != 0:
                response = response[start:end]
            response = response.replace('```json', '').replace('```', '')
            return response.strip()
        except Exception:
            return response

    def analyze_chunk(self, chunk: str, is_binary: bool = False) -> Dict:
        try:
            if is_binary:
                response = self.binary_analysis_chain.predict(strings_chunk=chunk)
            else:
                response = self.analysis_chain.predict(code_chunk=chunk)
                
            cleaned_response = self.clean_json_response(response)
            
            try:
                return json.loads(cleaned_response)
            except json.JSONDecodeError as je:
                return self._create_error_analysis("JSON parsing failed", str(je), is_binary)
                
        except Exception as e:
            return self._create_error_analysis("Analysis failed", str(e), is_binary)

    def _create_error_analysis(self, error_type: str, details: str, is_binary: bool = False) -> Dict:
        if is_binary:
            sections = [
                "suspicious_strings",
                "command_and_control_indicators",
                "anti_analysis_indicators",
                "network_related_strings",
                "file_system_indicators",
                "potential_malware_functionality"
            ]
        else:
            sections = [
                "code_obfuscation_techniques",
                "suspicious_api_calls",
                "anti_analysis_mechanisms",
                "network_communication_patterns",
                "file_system_operations",
                "potential_payload_analysis"
            ]
            
        return {
            "error": f"{error_type}: {details}",
            "summary": ["Analysis failed - " + error_type],
            "sections": {section: {
                "findings": ["Analysis failed"],
                "description": "Analysis failed due to technical error"
            } for section in sections}
        }

    def split_code_in_chunks(self, content: str, chunk_size: int = 12800) -> List[str]:
        return [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]

    def analyze_code(self, code_content: str) -> Dict:
        chunks = self.split_code_in_chunks(code_content)
        analyses = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk in enumerate(chunks, 1):
            status_text.text(f"Analyzing chunk {i}/{len(chunks)}...")
            analysis = self.analyze_chunk(chunk)
            analyses.append(analysis)
            progress_bar.progress(i/len(chunks))
        
        status_text.text("Analysis complete!")
        progress_bar.empty()
        
        return self.combine_analyses(analyses)
        

    def extract_strings_from_binary(self, binary_data: bytes, min_length: int = 10) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.exe') as temp_file:
            temp_file.write(binary_data)
            temp_file_path = temp_file.name

        try:
            try:
                # Use 'strings' with a length filter if available
                result = subprocess.run(['strings', '-n', str(min_length), temp_file_path],
                                        capture_output=True, text=True, check=True)
                extracted_strings = result.stdout
            except (subprocess.SubprocessError, FileNotFoundError):
                # Fallback: Extract printable ASCII strings manually
                extracted_strings = self._extract_strings_manually(binary_data, min_length)

            # Filter output (optional, e.g., extract only strings containing certain keywords)
            filtered_strings = "\n".join(
                line for line in extracted_strings.splitlines()
                if re.search(r'[a-zA-Z0-9_]', line)  # Ensure meaningful content
            )
            
            return " ".join(filtered_strings.splitlines())
        finally:
            try:
                os.unlink(temp_file_path)
            except:
                pass

    
    def _extract_strings_manually(self, binary_data: bytes, min_length: int = 4) -> str:
        strings = []
        current_string = ""
        
        for byte in binary_data:
            # Check if byte is printable ASCII
            if 32 <= byte <= 126:  # Printable ASCII range
                current_string += chr(byte)
            else:
                if len(current_string) >= min_length:
                    strings.append(current_string)
                current_string = ""
                
        # Add the last string if it meets the minimum length
        if len(current_string) >= min_length:
            strings.append(current_string)
            
        return "\n".join(strings)
    
    def analyze_binary(self, binary_data: bytes) -> Dict:
        strings_content = self.extract_strings_from_binary(binary_data)
        chunks = self.split_code_in_chunks(strings_content)
        analyses = []
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk in enumerate(chunks, 1):
            status_text.text(f"Analyzing binary chunk {i}/{len(chunks)}...")
            analysis = self.analyze_chunk(chunk, is_binary=True)
            analyses.append(analysis)
            progress_bar.progress(i/len(chunks))
        
        status_text.text("Binary analysis complete!")
        progress_bar.empty()
        
        return self.combine_analyses(analyses)

    def combine_analyses(self, analyses: List[Dict]) -> Dict:
        # Get all section keys from the first analysis
        if not analyses:
            return {
                "summary": ["No analysis results available"],
                "sections": {},
                "errors": ["No analyses performed"]
            }
            
        # Initialize combined structure with all possible sections
        combined = {
            "summary": set(),
            "sections": {},
            "errors": []
        }
        
        # Initialize sections based on first analysis
        if "sections" in analyses[0]:
            for section in analyses[0]["sections"]:
                combined["sections"][section] = {"findings": set(), "description": ""}
        
        for analysis in analyses:
            if "error" in analysis:
                combined["errors"].append(analysis["error"])
            
            if "summary" in analysis:
                combined["summary"].update(analysis["summary"])
            
            if "sections" in analysis:
                for section, content in analysis["sections"].items():
                    if section not in combined["sections"]:
                        combined["sections"][section] = {"findings": set(), "description": ""}
                    
                    if "findings" in content:
                        combined["sections"][section]["findings"].update(content["findings"])
                    
                    if content.get("description") and not combined["sections"][section]["description"]:
                        combined["sections"][section]["description"] = content["description"]

        result = {
            "summary": list(combined["summary"]),
            "sections": {},
            "errors": combined["errors"]
        }

        for section, content in combined["sections"].items():
            findings = list(content["findings"])
            if len(findings) > 1 and "Analysis failed" in findings:
                findings.remove("Analysis failed")
            
            result["sections"][section] = {
                "findings": findings,
                "description": content["description"] or "No significant findings in this category"
            }

        return result

    def create_analysis_report(self, analysis_results: Dict, title: str = "Security Analysis Report") -> str:
        document = Document()
        document.add_heading(title, 0)
        
        document.add_heading("Executive Summary", level=1)
        for summary_point in analysis_results.get("summary", ["No summary available"]):
            document.add_paragraph(summary_point, style='Body Text')
        
        for section_name, content in analysis_results.get("sections", {}).items():
            heading_text = section_name.replace('_', ' ').title()
            document.add_heading(heading_text, level=1)
            
            if content.get("description"):
                document.add_paragraph(content["description"], style='Body Text')
            
            if content.get("findings"):
                document.add_heading("Findings:", level=2)
                for finding in content["findings"]:
                    if finding != "None identified":
                        document.add_paragraph(f"• {finding}", style='List Bullet')
                    else:
                        document.add_paragraph("No specific issues identified in this category.", style='Body Text')
        
        if analysis_results.get("errors"):
            document.add_heading("Analysis Errors", level=1)
            for error in analysis_results["errors"]:
                document.add_paragraph(f"• {error}", style='List Bullet')
        
        report_filename = f"security_analysis_report_{os.getpid()}.docx"
        document.save(report_filename)
        return report_filename

    def get_chat_response(self, user_input: str) -> str:
        return self.conversation.predict(input=user_input+"Response should be short and crisp")

def display_analysis_results(analysis: Dict):
    st.header("Executive Summary")
    for summary_point in analysis.get("summary", ["No summary available"]):
        st.write(summary_point)
    st.divider()

    if analysis.get("errors"):
        st.error("Analysis Errors")
        for error in analysis["errors"]:
            st.write(f"• {error}")
        st.divider()

    for section_name, content in analysis.get("sections", {}).items():
        st.subheader(section_name.replace('_', ' ').title())
        
        if content.get("description"):
            st.write(content["description"])
        
        if content.get("findings"):
            st.write("Findings:")
            for finding in content["findings"]:
                st.write(f"• {finding}")
        st.divider()

def main():
    # Apply custom theme
    st.markdown(dark_purple_theme, unsafe_allow_html=True)
    
    # Initialize app
    if "app" not in st.session_state:
        st.session_state.app = SecurityAnalysisApp()
        st.session_state.messages = []
    
    # Sidebar with hover tabs
    with st.sidebar:
        tabs = on_hover_tabs(
            tabName=['Code Analyzer', 'GAMKERSGPT', 'About'],
            iconName=['code', 'chat', 'info'],
            styles={
                        'navtab': {
                            'background-color': 'black',
                            'color': '#9D4EDD',
                            'font-size': '16px',
                            'transition': '.3s',
                            'white-space': 'nowrap',
                            'text-transform': 'uppercase'
                        },
                        'tabOptionsStyle': {
                            ':hover': {'color': '#1A1A1A', 'background-color': 'black'}
                        },
                        'iconStyle': {
                            'position': 'fixed',
                            'left': '7.5px',
                            'text-align': 'left'
                        },
                        'tabStyle': {
                            'list-style-type': 'none',
                            'margin-bottom': '30px',
                            'padding-left': '30px'
                        }
                    }
        )

    if tabs == 'About':
        about_section()
    elif tabs == 'Code Analyzer':
        st.title("GAMKERS Security Analyzer")
        tab1, tab2, tab3 = st.tabs(["📝 Paste Code", "📁 Upload Source File", "💾 Upload Binary"])

        with tab1:
            code_input = st.text_area("Paste your code here:", height=300)
            if st.button("🔍 Analyze Code", key="analyze_pasted") and code_input:
                with st.spinner("🔄 Analyzing code..."):
                    analysis_results = st.session_state.app.analyze_code(code_input)
                    display_analysis_results(analysis_results)
                    report_filename = st.session_state.app.create_analysis_report(analysis_results)
                    with open(report_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Analysis Report",
                            data=file,
                            file_name=report_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(report_filename)

        with tab2:
            uploaded_file = st.file_uploader("Choose a source code file", type=['py', 'js', 'java', 'cpp', 'cs', 'php', 'rb'], key="code_file")
            if st.button("🔍 Analyze Source File", key="analyze_source") and uploaded_file:
                with st.spinner("🔄 Analyzing source file..."):
                    code_content = uploaded_file.read().decode()
                    analysis_results = st.session_state.app.analyze_code(code_content)
                    display_analysis_results(analysis_results)
                    report_filename = st.session_state.app.create_analysis_report(analysis_results)
                    with open(report_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Analysis Report",
                            data=file,
                            file_name=report_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(report_filename)
        
        with tab3:
            st.write("Upload a binary file (.exe) for security analysis")
            uploaded_binary = st.file_uploader("Choose a binary file", type=['exe'], key="binary_file")
            
            if uploaded_binary:
                st.info("Binary analysis will extract strings from the executable and analyze them for security indicators.")
            
            if st.button("🔍 Analyze Binary", key="analyze_binary") and uploaded_binary:
                with st.spinner("🔄 Extracting strings and analyzing binary..."):
                    binary_data = uploaded_binary.read()
                    
                    # Create expandable section to show extracted strings
                    with st.expander("Extracted Strings Preview"):
                        extracted_strings = st.session_state.app.extract_strings_from_binary(binary_data)
                        st.text_area("Strings from binary", value=extracted_strings[:10000] + 
                                    ("\n\n[Truncated...]" if len(extracted_strings) > 10000 else ""), 
                                    height=300, disabled=True)
                    
                    analysis_results = st.session_state.app.analyze_binary(binary_data)
                    st.subheader("Binary Analysis Results")
                    display_analysis_results(analysis_results)
                    
                    report_filename = st.session_state.app.create_analysis_report(
                        analysis_results, 
                        title=f"Binary Security Analysis Report - {uploaded_binary.name}"
                    )
                    
                    with open(report_filename, "rb") as file:
                        st.download_button(
                            label="📥 Download Binary Analysis Report",
                            data=file,
                            file_name=report_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(report_filename)

    elif tabs == 'GAMKERSGPT':
        st.title("GAMKERSGPT - DeepSeek-R1")
        st.markdown(
            """
            <style>
            /* Style for chat messages */
            [data-testid="stChatMessage"] {
                background-color: #000000 !important;
                color: #ffffff !important;
                padding: 10px;
                border-radius: 8px;
            }
            /* Style for chat input */
            [data-testid="stChatInput"] {
                background-color: #000000 !important;
                color: #ffffff !important;
                border: 1px solid #333333;
                border-radius: 8px;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.markdown(
            """
            <style>
            /* Target the container that holds the chat input (stBottom) */
            [data-testid="stBottom"] {
                background-color: #000000 !important;  /* Set the background to black */
                padding: 10px !important;              /* Optional: adjust padding */
                border-radius: 8px !important;         /* Optional: add rounded corners */
            }
            </style>
            """,
            unsafe_allow_html=True
        )

        # Display chatbot information before the chat input
        st.markdown("### Welcome to GAMKERSGPT!")
        st.write(
            """
            GAMKERSGPT is your cyber security assistant designed to help you with your security-related queries.
            Ask about network vulnerabilities, best practices, threat analysis, and more.
            Please note that while GAMKERSGPT aims to provide accurate information, it is for informational purposes only.
            """
        )
        st.info("Tip: Be as specific as possible with your questions to get the most relevant answers.")
        
        # Display existing chat messages
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Chat input for new messages
        if prompt := st.chat_input("Ask your cyber security question..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                with st.spinner("🤔 Thinking..."):
                    response = st.session_state.app.get_chat_response(prompt)
                    response=response.split("</think>")[-1]
                    st.markdown(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})


if __name__ == "__main__":
    main()
